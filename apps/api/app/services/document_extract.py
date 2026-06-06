"""Extract plain text from an uploaded resume / job-description document.

Supports the formats people actually drag in: PDF (``pypdf``), Word ``.docx``
(``python-docx``) and plain ``.txt`` / ``.md``. Parsing is CPU-bound and the
libraries are synchronous, so callers should run :func:`extract_text` in a
thread (e.g. ``anyio.to_thread.run_sync``) to avoid blocking the event loop.
"""

from __future__ import annotations

import io
from enum import StrEnum

import structlog

log = structlog.get_logger(__name__)

# Guard rails so a hostile or accidental upload can't exhaust memory/CPU.
MAX_FILE_BYTES = 8 * 1024 * 1024  # 8 MB
MAX_TEXT_CHARS = 60_000


class DocumentError(ValueError):
    """Raised when a document can't be read or yields no usable text."""


class DocumentKind(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    TEXT = "text"


_PDF_TYPES = {"application/pdf", "application/x-pdf"}
_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_TEXT_TYPES = {"text/plain", "text/markdown"}


def detect_kind(filename: str | None, content_type: str | None) -> DocumentKind:
    """Resolve the document kind from MIME type first, then file extension."""
    ctype = (content_type or "").split(";")[0].strip().lower()
    if ctype in _PDF_TYPES:
        return DocumentKind.PDF
    if ctype in _DOCX_TYPES:
        return DocumentKind.DOCX
    if ctype in _TEXT_TYPES:
        return DocumentKind.TEXT

    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return DocumentKind.PDF
    if name.endswith(".docx"):
        return DocumentKind.DOCX
    if name.endswith((".txt", ".md", ".markdown")):
        return DocumentKind.TEXT

    raise DocumentError(
        "Unsupported file type. Drop a PDF, Word .docx, or .txt file."
    )


def extract_text(
    data: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> tuple[str, DocumentKind]:
    """Return ``(clean_text, kind)`` extracted from the raw upload bytes.

    Raises :class:`DocumentError` on unsupported types, oversized files, or
    documents that contain no extractable text (e.g. a scanned/image-only PDF).
    """
    if not data:
        raise DocumentError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise DocumentError(
            f"File is too large ({len(data) // 1024} KB). Max is {MAX_FILE_BYTES // (1024 * 1024)} MB."
        )

    kind = detect_kind(filename, content_type)
    if kind is DocumentKind.PDF:
        text = _extract_pdf(data)
    elif kind is DocumentKind.DOCX:
        text = _extract_docx(data)
    else:
        text = _extract_plain(data)

    cleaned = _normalize(text)
    if not cleaned:
        raise DocumentError(
            "No readable text was found. If this is a scanned PDF, paste the text manually."
        )
    return cleaned[:MAX_TEXT_CHARS], kind


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise DocumentError("This PDF appears to be corrupt or unreadable.") from exc

    if reader.is_encrypted:
        # Many resumes are exported with an empty owner password; try that once.
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 - any failure means we can't read it
            raise DocumentError("This PDF is password-protected.") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - skip a bad page, keep the rest
            log.warning("pdf_page_extract_failed", error=str(exc))
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(io.BytesIO(data))
    except PackageNotFoundError as exc:
        raise DocumentError("This doesn't look like a valid .docx file.") from exc

    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize(text: str) -> str:
    """Collapse runaway whitespace while preserving paragraph/line structure."""
    lines = [" ".join(line.split()) for line in text.replace("\r\n", "\n").split("\n")]
    out: list[str] = []
    blank_run = 0
    for line in lines:
        if line:
            out.append(line)
            blank_run = 0
        else:
            blank_run += 1
            if blank_run <= 1:  # keep single blank lines, drop the rest
                out.append("")
    return "\n".join(out).strip()
